# -*- coding: utf-8 -*-
# made by lck

# 环境说明：(使用须知！！！！！！很重要，必看！！！！！）

# 1. 运行依赖：需安装 Python 3.6 及以上版本（推荐 3.8+，兼容性更优）

# 2. 必要库：需提前安装处理 Word 和 Excel 的专用库，安装命令：
# pip install openpyxl python-docx

# 若安装速度慢，可使用清华大学镜像：
# pip install openpyxl python-docx -i https://pypi.tuna.tsinghua.edu.cn/simple

# 3. 系统兼容性：支持 Windows、macOS、Linux 系统，文件路径需按系统格式填写：
# - Windows 路径示例：E:\eg\文件夹\eg.docx
# - macOS/Linux 路径示例：/Users/用户名/eg/文件夹/防火２有支撑版.xlsx

# 4. 注意事项：
# - Word 源文件需为 .docx 格式，数据需存储在含 “测点 1”“平均值” 关键词的表格中（程序仅识别此类表格）
# - Excel 模板统一使用 “防火２有支撑版.xlsx”，未使用的工作表（如“支撑”）会自动清理
# - 运行时请关闭目标 Word 和 Excel 文件，避免文件占用导致读写失败或数据损坏
# - 程序会自动生成 “汇总原始记录.docx” 并存于 Word 同目录，用于数据核对
# - 支持 “钢柱”“钢梁”“支撑” 分类，未识别构件自动归为 “其他” 类，共用钢柱模板格式
# - 生成的 Excel 报告自动命名为 “The Unification_报告版.xlsx”，同名文件会自动加序号（如 “The Unification_报告版 (1).xlsx”）
# - “μ” 字符自动适配 Times New Roman 字体；仪器型号按平均值自动识别（<10→23-90，≥10→24-57）
# - 日期分桶模式支持规则重叠处理，默认按 “后面的天” 优先，未分配数据可通过输入 “a” 并入最后一天

# === The Unification ===


from pathlib import Path
import re, copy, math, warnings, sys, os, unicodedata, ctypes
from collections import defaultdict
from datetime import datetime
from docx import Document
from docx.shared import RGBColor, Pt
from openpyxl.styles import Font, Alignment

# made by lck, an intern of this company in 2025 summer

warnings.filterwarnings("ignore", category=SyntaxWarning)

TITLE = "The Unification"
VERSION = "v 4.2.2"
AUTHOR = "LCK"

# ===== 默认路径 =====
WORD_SRC_DEFAULT = Path(r"D:\eg\eg.docx")
XLSX_WITH_SUPPORT_DEFAULT = Path(r"E:\公司尝试\防火原始文件\防火excel模板.xlsx")
DEFAULT_FONT_PT = 9

# 每页 5 组、每组 5 行、每行 8 读数+平均值
PER_LINE_PER_BLOCK = 5
BLOCKS_PER_SHEET = 5

# 本次运行只提示一次
_hint_shown = False

# 打印顺序：可自行调整位置
CATEGORY_ORDER = ["钢柱", "钢梁", "支撑", "其他"]

# 支撑分桶策略："number"=按编号，"floor"=按楼层；仅本次运行生效
support_bucket_strategy = None


# === 通用输入封装 ===


def enable_ansi():
    if os.name != "nt":
        return True
    k32 = ctypes.windll.kernel32
    h = k32.GetStdHandle(-11)  # STD_OUTPUT_HANDLE
    mode = ctypes.c_uint32()
    if not k32.GetConsoleMode(h, ctypes.byref(mode)):
        return False
    return bool(k32.SetConsoleMode(h, mode.value | 0x0004))  # ENABLE_VIRTUAL_TERMINAL_PROCESSING


enable_ansi()

# 颜色：暗灰（bright black）+ 微弱（dim）
DIM = "\x1b[2m"
GRAY = "\x1b[90m"
RESET = "\x1b[0m"


def dark_hint(text: str) -> str:
    """
    输出极深灰提示（几乎黑）。优先用 truecolor；否则退回 256 色 232。
    """
    # truecolor（24-bit）
    try:
        return f"\x1b[2m\x1b[38;2;12;12;12m{text}{RESET}"  # (12,12,12) 比 (18,18,18) 更贴近黑
    except Exception:
        # 256 色兜底：232 是最暗的灰阶
        return f"\x1b[2m\x1b[38;5;232m{text}{RESET}"


class BackStep(Exception):
    """用户输入 q 请求返回上一步。"""
    pass


class AbortToPath(Exception):
    """用户主动中断当前模式并返回路径输入。"""
    pass


def ask(prompt: str, allow_empty: bool = True, lower: bool = False) -> str:
    """统一的控制台输入函数。

    参数:
        prompt: 提示字符串。
        allow_empty: 是否允许空输入；False 时会重复询问。
        lower: 返回值是否小写化。

    返回:
        用户输入的字符串（可小写化）。

    特殊:
        输入 ``q`` 将触发 :class:`BackStep` 异常。
        仅识别小写 ``q``，大写 ``Q`` 在此阶段视为普通字符。
    """
    while True:
        raw = input(f"{prompt}\n→ ").strip()
        if raw == "q":
            raise BackStep()
        if not allow_empty and raw == "":
            continue
        return raw.lower() if lower else raw


def show_help_browser():
    """帮助浏览器包装。"""
    tutorial_browser()


def show_easter_egg():
    """Easter egg message for curious users."""
    print("\n🎉这是一个小彩蛋，致正在北京漂泊的你："

          """\n          嘿，今天过得怎样？
          
          有没有如愿多睡一会懒觉，有没有觉得自己比昨天更好
          
          我想听听你今天的小事——
          老板有没有临下班给你丢个“顺手看看”？
          外卖是不是还是那家麻辣烫，你点“微辣”结果还是上头？
          回到合租房，你是不是又轻轻关门，怕惊醒陌生的梦？
          
          别急着坚强，先放松一会儿。咱慢慢说。
          
          有时候我也会想：我们到底在赶什么？
          通勤像回合制游戏，卡点打卡，换乘升级；
          朋友圈像展览，大家都把光亮挂在墙上，阴影藏在鞋盒里。
          你说你累，我懂——不是“做事”的累，是“证明自己”的累。
          
          但是你知道吗，我喜欢看你认真时那个表情：
          眉心轻轻拧一下，像在和困难开私聊；
          打完一行代码、写完一段文案、对齐一张表格，
          你会悄悄点一下保存，像给自己递水。
          
          北京没有义务温柔，但我们可以对彼此温柔。
          你讲，我听；我讲，你也可以打断我。
          我们不解决所有问题，只把今晚的叹气放下三分之一就行。
          
          如果你问“值得吗？”
          我也会反问你：“哪一刻让你觉得还想再试一次？”
          是凌晨的页面通过了，是邮件里突然多了个“已阅”，
          还是朋友说了一句“有你真好”？
          这些微小的亮，它们不大，但够我们往前挪半步。
          
          我不劝你乐观，也不催你振作。
          我只想把这句话放在这里，像把外套搭在你肩上：
          
          我们可以慢一点，但别把自己弄丢。
          
          等你想继续聊，我还在。
          在五环的风里，在灯没关的屏幕前，在你回消息的那个“嗯”字后面。
          
          晚安，先把背放松，再把心放下。
          明天见，我们接着说。
          LCK
          """)

    input("按回车即可返回")


def ask_path() -> str | None:
    """顶层路径输入。

    返回 ``None`` 表示用户查看帮助后继续；
    返回 ``"__QUIT__"`` 表示用户请求退出程序；
    其他返回值为用户输入的路径字符串。
    """
    raw = input("📂 请输入 Word 源路径（eg：D:\某防火.docx）\n→ ").strip()
    if unicodedata.normalize('NFKC', raw).lower() == "k":
        show_easter_egg()
        return None
    if raw == "help":
        show_help_browser()
        return None
    if raw == "Q":
        return "__QUIT__"
    return raw


def is_valid_path(p: str) -> bool:
    """简单校验路径是否存在。"""
    path_obj = Path(p.strip('"'))
    return path_obj.exists() and path_obj.is_file()


# ---- 文件占用友好提示封装 ----
class FileInUse(Exception):
    pass


def _is_in_use_error(e: Exception) -> bool:
    # Windows 常见：WinError 32（共享冲突），或 PermissionError 13
    msg = str(e).lower()
    code32 = getattr(e, "winerror", None) == 32
    perm13 = isinstance(e, PermissionError)
    hit_msg = ("being used by another process" in msg or
               "used by another process" in msg or
               "permission denied" in msg)
    return bool(code32 or perm13 or hit_msg)


def load_workbook_safe(path, **kw):
    from openpyxl import load_workbook
    try:
        return load_workbook(path, **kw)
    except Exception as e:
        if _is_in_use_error(e):
            raise FileInUse(f"Excel 模板/文件被占用：{path}") from e
        raise


def save_workbook_safe(wb, path):
    try:
        wb.save(path)
    except Exception as e:
        if _is_in_use_error(e):
            raise FileInUse(f"无法保存 Excel（被占用）：{path}") from e
        raise


def save_docx_safe(doc, path):
    try:
        doc.save(str(path))
    except Exception as e:
        if _is_in_use_error(e):
            raise FileInUse(f"无法保存 Word（被占用）：{path}") from e
        raise


# ===== Word 汇总生成 =====
NEED_COLS = 11
MIN_ROWS_EACH = 5
PLACEHOLDER = "/"
digit_re = re.compile(r"\d")
HEADER = [
    "序号", "构件名称及部位",
    "测点1 读数1", "测点1 读数2",
    "测点2 读数1", "测点2 读数2",
    "测点3 读数1", "测点3 读数2",
    "测点4 读数1", "测点4 读数2",
    "涂层厚度平均值"
]


def ensure_cells(row, need=NEED_COLS):
    """
    确保表格行包含足够的单元格，不足时自动补充空白单元格。

    通过复制首个单元格的格式创建空白单元格，避免因原始表格列数不足导致数据提取失败，保障数据结构完整性。

    Args:
        row: Word表格行对象（docx.table.Row）
        need: 需要的最小列数，默认11列（与汇总表列数一致）
    """
    while len(row.cells) < need:
        tc = copy.deepcopy(row.cells[0]._tc)  # noqa
        for t in tc.xpath('.//*[local-name()="t"]'): t.text = ''
        row._tr.append(tc)  # noqa


def color_row_red(row):
    """
    将表格行的文字颜色设置为红色，用于表头高亮显示。

    通过遍历行内所有单元格和段落，统一设置文字颜色为红色，增强汇总表中表头与数据行的区分度。

    Args:
        row: Word表格行对象（docx.table.Row）
    """
    for c in row.cells:
        for p in c.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 0, 0)


def is_data_table(tbl):
    """
    判断Word表格是否为有效数据表格（含测点和平均值信息）。

    通过检查表格前3行是否同时包含“测点1”和“平均值”关键词，筛选出实际存储检测数据的表格，排除说明性表格。

    Args:
        tbl: Word表格对象（docx.table.Table）
    Returns:
        bool: 是有效数据表格则返回True，否则返回False
    """
    first_three = " ".join(c.text for r in tbl.rows[:3] for c in r.cells)
    return "测点1" in first_three and "平均值" in first_three


def detect_layout(tbl):
    """
    检测数据表格的列布局，确定测点列、平均值列位置及是否为钢梁表格。

    定位含“测点1”的表头行，提取测点列索引和平均值列索引；通过测点列数量判断是否为钢梁表格（钢梁含3个测点）。

    Args:
        tbl: Word表格对象（docx.table.Table）
    Returns:
        tuple: 包含三个元素的元组，分别为：
            - 测点列索引列表（list[int]）
            - 平均值列索引（int）
            - 是否为钢梁表格（bool，钢梁表格返回True）
    """
    hdr = next(r for r in tbl.rows if "测点1" in "".join(c.text for c in r.cells))
    col_vals, col_avg = [], None
    for i, t in enumerate(hdr.cells):
        txt = (t.text or "").strip()
        m = re.match(r"测点(\d+)", txt)
        if m:
            col_vals.append(i)
        elif "平均值" in txt and "所有" not in txt:
            col_avg = i
    is_beam = len(col_vals) == 3  # 梁 3 组，柱/支撑 4 组
    return col_vals, col_avg, is_beam


def extract_rows_with_progress(tbl, ti: int, T: int):  # noqa
    """
    从数据表格提取行数据，带实时进度提示。

    按表头布局提取构件名称、测点值和平均值，对钢梁表格自动补充第4个测点（用“/”占位）；通过控制台实时显示提取进度（按行计算）。

    Args:
        tbl: Word表格对象（docx.table.Table）
        ti: 当前表格在总表格中的序号（从1开始）
        T: 需处理的总表格数量
    Returns:
        list[dict]: 提取的数据行列表，每个元素为包含以下键的字典：
            - name: 构件名称（str）
            - vals: 测点值列表（list[str]）
            - avg: 平均值（str）
            - is_hdr: 是否为表头行（bool）
    """
    col_vals, col_avg, is_beam = detect_layout(tbl)
    rows, last_comp, last_avg = [], None, ""
    buffer = []

    total = len(tbl.rows)
    last_flush = -1

    for ridx, r in enumerate(tbl.rows):
        if ridx // 20 != last_flush:
            last_flush = ridx // 20
            pct = int((ridx + 1) * 100 / max(1, total))
            sys.stdout.write(f"\r📝 读取 Word：表 {ti}/{T}（{pct}%）")
            sys.stdout.flush()

        line = " ".join(c.text for c in r.cells)

        if "测点1" in line:
            if buffer:
                rows.extend(buffer);
                buffer.clear()  # noqa
            meas_titles = [f"测点{i + 1}" for i in range(len(col_vals))]
            if is_beam: meas_titles.append("测点4")  # 梁补第4列标题
            rows.append({"name": "", "vals": meas_titles, "avg": "平均值", "is_hdr": True})
            continue

        if not digit_re.search(line):
            continue

        comp = r.cells[1].text.strip()
        vals = [r.cells[i].text.strip() for i in col_vals]
        if is_beam and len(vals) == 3: vals.append("/")

        raw_avg = r.cells[col_avg].text.replace("\n", "").strip()
        avg = raw_avg or last_avg or "/"
        last_avg = avg if raw_avg else last_avg

        buffer.append({"name": comp if comp != last_comp else "",
                       "vals": vals, "avg": avg, "is_hdr": False})
        last_comp = comp

    rows.extend(buffer)
    sys.stdout.write(f"\r📝 读取 Word：表 {ti}/{T}（100%）\n");
    sys.stdout.flush()
    return rows


def build_summary_doc_with_progress(rows):
    """
     生成Word汇总表，带实时进度提示。

     将提取的数据行整理为规范表格，表头标红；不足行数用占位符补充，统一字体大小；通过控制台显示组装进度。

     Args:
         rows: 提取的数据行列表（extract_rows_with_progress返回结果）
     Returns:
         Document: 生成的Word汇总表文档对象（docx.document.Document）
     """
    doc = Document()
    tbl = doc.add_table(rows=1, cols=NEED_COLS)
    tbl.style = "Table Grid"
    for i, t in enumerate(HEADER):
        tbl.rows[0].cells[i].text = t
    color_row_red(tbl.rows[0])

    serial, last_comp, buffer = 1, None, []
    total = len(rows)
    step = max(50, total // 100)

    def flush():
        nonlocal serial, buffer
        miss = max(0, MIN_ROWS_EACH - len(buffer))
        for _ in range(miss):
            q = tbl.add_row();
            ensure_cells(q)
            for z in range(2, 10): q.cells[z].text = PLACEHOLDER
            q.cells[10].text = PLACEHOLDER
        serial += 1;
        buffer.clear()

    for i, it in enumerate(rows, start=1):
        if i % step == 0 or i == total:
            pct = int(i * 100 / max(1, total))
            sys.stdout.write(f"\r📦 组装汇总：{i}/{total}（{pct}%）")
            sys.stdout.flush()

        if it["is_hdr"] and buffer: flush()

        raw_name = (it.get("name") or "").strip()
        comp = raw_name or last_comp or ""

        if last_comp and comp and comp != last_comp:
            flush();
            last_comp = None

        if it.get("is_hdr"):
            r = tbl.add_row();
            ensure_cells(r);
            color_row_red(r)
            r.cells[1].text = "构件名称及部位" if not raw_name else raw_name
            for k, v in enumerate(it["vals"]):
                c = 2 + k * 2
                r.cells[c].text = v
            r.cells[10].text = it["avg"]
            last_comp = comp
            continue

        r = tbl.add_row();
        ensure_cells(r);
        buffer.append(r)
        first = (last_comp is None) or (comp and comp != last_comp)
        if first:
            r.cells[0].text = str(serial)
            r.cells[1].text = raw_name
            last_comp = comp
        for k, v in enumerate(it["vals"]):
            c = 2 + k * 2
            r.cells[c].text = v
            r.cells[c + 1].text = v
        r.cells[10].text = it["avg"]

    flush()
    sys.stdout.write("\n");
    sys.stdout.flush()
    return doc


def set_doc_font_progress(doc, pt=DEFAULT_FONT_PT):
    """
    统一Word文档中所有文字的字体大小，带实时进度提示。

    遍历文档中的所有段落和表格单元格，将字体大小设置为指定磅数（默认9pt）；通过控制台显示字体设置进度。

    Args:
        doc: Word文档对象（docx.document.Document）
        pt: 字体大小（磅），默认9pt
    """
    cell_pars = 0
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                cell_pars += len(c.paragraphs)
    total = len(doc.paragraphs) + cell_pars
    done = 0
    step = max(200, total // 100)

    for p in doc.paragraphs:
        for run in p.runs: run.font.size = Pt(pt)
        done += 1
        if done % step == 0 or done == total:
            pct = int(done * 100 / max(1, total))
            sys.stdout.write(f"\r🖋 统一字体：{done}/{total}（{pct}%）");
            sys.stdout.flush()

    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    for run in p.runs: run.font.size = Pt(pt)
                    done += 1
                    if done % step == 0 or done == total:
                        pct = int(done * 100 / max(1, total))
                        sys.stdout.write(f"\r🖋 统一字体：{done}/{total}（{pct}%）");
                        sys.stdout.flush()
    sys.stdout.write("\n");
    sys.stdout.flush()


# ===== rows → groups（8读数+平均值）=====
def groups_from_your_rows(rows_all_tables):
    """
    将提取的原始数据行转换为按构件分组的结构化数据。

    按构件名称分组，将每组数据整理为规范格式（8个读数+1个平均值），自动用“/”补齐不足的读数。

    Args:
        rows_all_tables: 所有表格提取的原始数据行列表（extract_rows_with_progress返回结果）
    Returns:
        list[dict]: 构件数据组列表，每个元素为包含以下键的字典：
            - name: 构件名称（str）
            - data: 数据行列表，每行包含8个读数和1个平均值（list[list[str]]）
    """
    groups = [];
    cur = None
    for it in rows_all_tables:
        if it.get("is_hdr"): continue
        name = (it.get("name") or "").strip()
        if name:
            if cur and cur["data"]: groups.append(cur)  # noqa
            cur = {"name": name, "data": []}
        if not cur: continue
        vals8 = []
        for v in it["vals"]:
            v = (v or "/").strip() or "/"
            vals8.extend([v, v])
        while len(vals8) < 8: vals8.append("/")
        avg = (it.get("avg") or "/").strip() or "/"
        cur["data"].append(vals8[:8] + [avg])  # noqa
    if cur and cur["data"]: groups.append(cur)
    return groups


# ===== 分类 / 规则 =====
CATEGORY_SYNONYMS = {
    "支撑": ["支撑", "WZ", "ZC", "支架", "斜撑", "撑杆"],
    "钢柱": ["钢柱", "柱", "GZ", "框架柱", "立柱", "H柱"],
    "钢梁": ["钢梁", "梁", "GL", "连系梁", "檩条", "楼梯梁", "平台梁", "屋架梁"],
}


def kind_of(name: str) -> str:
    """
    根据构件名称判断类型（钢柱/钢梁/支撑/其他）。

    基于预设的同义词表匹配构件名称中的关键词（如“钢柱”或“GZ”对应钢柱），未匹配到关键词的构件归为“其他”类。

    Args:
        name: 构件名称字符串（str）
    Returns:
        str: 构件类型，可能为"钢柱"、"钢梁"、"支撑"或"其他"
    """
    s_up = name.upper()
    for cat, words in CATEGORY_SYNONYMS.items():
        for w in words:
            if w.isascii():
                if w.upper() in s_up:
                    return cat
            else:
                if w in name:
                    return cat
    return "其他"  # 未识别 → 其他


def floor_of(name: str) -> int:
    """
    从构件名称中提取楼层号，特殊楼层用固定大数值标记。

    支持多种格式（如“5层”“F5”“L5”），屋面/顶层标记为10⁶，机房层标记为10⁶-1，地下室/负楼层返回0。

    Args:
        name: 构件名称字符串（str）
    Returns:
        int: 提取的楼层号（特殊楼层用10⁶级数值，无楼层信息返回0）
    """
    s = name.replace("－", "-").replace("—", "-").replace("–", "-")
    sl = s.lower()
    if re.search(r"(?:屋面|屋顶|顶\s*层)", s) or re.search(r"\b(?:wm|dc)", sl): return 10 ** 6  # noqa
    if "机房层" in s or re.search(r"\bjf", sl): return 10 ** 6 - 1
    m = re.search(r"(?i)[FL]\s*(\d+)", s)
    if m: return int(m.group(1))
    m = re.search(r"(?i)(\d+)\s*[FL]", s)
    if m: return int(m.group(1))
    m = re.search(r"(\d+)\s*[层樓楼]", s)
    if m: return int(m.group(1))
    if re.search(r"(?i)\bB\s*\d+\b|负\s*\d+\s*层?", s): return 0
    return 0


def _floor_label_from_name(name: str) -> str:
    """根据名称提取楼层标签，如"5F"、"B2"、"屋面"等。"""
    s = (name or "").replace("－", "-").replace("—", "-").replace("–", "-")
    sl = s.lower()
    if re.search(r"屋面|顶层", s) or re.search(r"\b(?:wm|dc)", sl):
        return "屋面"
    if "机房层" in s or re.search(r"\bjf", sl):
        return "机房层"
    m = re.search(r"(?i)B\s*(\d+)", s)
    if m:
        return f"B{int(m.group(1))}"
    m = re.search(r"(\d+)\s*[Ff层樓楼]?", s)
    if m:
        return f"{int(m.group(1))}F"
    return "F?"


def _floor_sort_key_by_label(label: str):
    """生成楼层标签的排序键。"""
    m = re.fullmatch(r"B(\d+)", label)
    if m:
        return (0, -int(m.group(1)))
    m = re.fullmatch(r"(\d+)F", label)
    if m:
        return (1, int(m.group(1)))
    if label == "机房层":
        return (2, 0)
    if label == "屋面":
        return (3, 0)
    return (4, 0)


def segment_index(floor: int, breaks: list[int]) -> int:
    """
    根据楼层断点计算当前楼层所属的分段索引，用于楼层分页逻辑。

    遍历断点列表，返回当前楼层首次小于等于断点值的索引；若大于所有断点，返回断点列表长度（最后一段）。

    Args:
        floor: 楼层号（int）
        breaks: 楼层断点列表（升序排列，list[int]）
    Returns:
        int: 分段索引（从0开始）
    """
    for i, b in enumerate(breaks):
        if floor <= b: return i
    return len(breaks)


def expand_blocks(groups, block_size=PER_LINE_PER_BLOCK):
    """
    将构件数据组拆分为固定大小的数据块（默认5行/块），不足行数用“/”补齐。

    按指定块大小（默认5行）拆分每组数据，确保每个块结构统一，适配Excel模板中“每组数据占5行”的格式要求。

    Args:
        groups: 构件数据组列表（groups_from_your_rows返回结果）
        block_size: 每个数据块的行数，默认5行
    Returns:
        list[dict]: 数据块列表，每个元素为包含以下键的字典：
            - name: 构件名称（str）
            - data: 5行数据（每行9列，list[list[str]]）
    """
    blocks = []
    for g in groups:
        rows = list(g["data"])
        for k in range(0, len(rows), block_size):
            sub = rows[k:k + block_size]
            while len(sub) < block_size: sub.append(['/'] * 9)
            blocks.append({"name": g["name"], "data": sub})
    return blocks


# ===== Excel sheet 复制与设置 =====
def clone_sheet_keep_print(wb, tpl_name: str, title: str):
    """
    复制Excel工作表并保留打印格式和视图设置，确保新表与模板格式一致。

    复制内容包括视图（缩放、冻结窗格）、打印区域、页面设置（方向、纸张大小）、页边距、行列宽等，保障格式统一性。

    Args:
        wb: Excel工作簿对象（openpyxl.workbook.Workbook）
        tpl_name: 模板工作表名称（str）
        title: 新工作表名称（str）
    Returns:
        openpyxl.worksheet.worksheet.Worksheet: 新复制的工作表对象
    """
    tpl = wb[tpl_name]
    ws = wb.copy_worksheet(tpl)
    ws.title = title
    ws.sheet_view.view = "pageBreakPreview"
    try:
        ws.freeze_panes = tpl.freeze_panes
    except:
        pass
    try:
        ws.print_area = tpl.print_area
    except:
        pass
    try:
        ws.print_titles = tpl.print_titles
    except:
        pass
    for attr in (
    "orientation", "paperSize", "fitToWidth", "fitToHeight", "scale", "firstPageNumber", "useFirstPageNumber"):
        try:
            setattr(ws.page_setup, attr, getattr(tpl.page_setup, attr))
        except:
            pass
    for attr in ("left", "right", "top", "bottom", "header", "footer"):
        try:
            setattr(ws.page_margins, attr, getattr(tpl.page_margins, attr))
        except:
            pass
    for col, dim in tpl.column_dimensions.items():
        if dim.width is not None:
            ws.column_dimensions[col].width = dim.width
    for row, dim in tpl.row_dimensions.items():
        if dim.height is not None:
            ws.row_dimensions[row].height = dim.height
    return ws


def ensure_total_pages(wb, base: str, total_needed: int):
    """
    确保Excel中有足够的指定类型工作表，不足时自动从基础表复制补充。

    筛选并排序已有同类型工作表，若数量不足，以基础表为模板复制新表并按序号命名（如“钢柱（2）”）。

    Args:
        wb: Excel工作簿对象（openpyxl.workbook.Workbook）
        base: 基础工作表名称（如"钢柱"，str）
        total_needed: 需要的工作表总数（int）
    Returns:
        list[str]: 排序后的工作表名称列表
    """
    names = [s for s in wb.sheetnames if s == base or re.match(rf'^{re.escape(base)}（\d+）$', s)]
    names = sorted(names, key=lambda n: 0 if n == base else int(re.findall(r'（(\d+)）', n)[0]))
    have = len(names)
    start = have + 1
    for _ in range(max(0, total_needed - have)):
        nm = f"{base}（{start}）"
        clone_sheet_keep_print(wb, base, nm)
        names.append(nm);
        start += 1
    return names


def ensure_total_pages_from(wb, tpl_name: str, new_base: str, total_needed: int):
    """
    为“其他”类构件确保足够的工作表，复用已有表或从指定模板复制。

    适用于无专用模板的类别，筛选已有同类型工作表，不足时从指定模板（如“钢柱”）复制新表并命名。

    Args:
        wb: Excel工作簿对象（openpyxl.workbook.Workbook）
        tpl_name: 模板工作表名称（如"钢柱"，str）
        new_base: 新类别基础名称（如"其他"，str）
        total_needed: 需要的工作表总数（int）
    Returns:
        list[str]: 排序后的工作表名称列表
    """
    # 复用已有“其他（n）”等；不足则从 tpl_name 复制
    names = [s for s in wb.sheetnames if s == new_base or re.match(rf'^{re.escape(new_base)}（\d+）$', s)]
    names = sorted(names, key=lambda n: 0 if n == new_base else int(re.findall(r'（(\d+)）', n)[0]))
    have = len(names)
    start = have + 1
    for _ in range(max(0, total_needed - have)):
        nm = f"{new_base}（{start}）" if start > 1 else new_base
        clone_sheet_keep_print(wb, tpl_name, nm)
        if nm not in names: names.append(nm)
        start += 1
    return names


def enforce_mu_font(wb):
    """
    遍历Excel所有单元格，将含“μ”字符的单元格字体强制设为Times New Roman。

    解决“μ”符号在部分字体下显示异常的问题，保留原字体的大小、加粗等其他属性。

    Args:
        wb: Excel工作簿对象（openpyxl.workbook.Workbook）
    """
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                v = cell.value
                if isinstance(v, str) and "μ" in v:
                    f = cell.font
                    cell.font = Font(
                        name="Times New Roman",
                        sz=f.sz, bold=f.bold, italic=f.italic, vertAlign=f.vertAlign,
                        underline=f.underline, strike=f.strike, color=f.color,
                        charset=f.charset, scheme=f.scheme, outline=f.outline
                    )


# ===== 数据区定位 / 写入 =====
def detect_anchors(ws):
    """
    检测Excel工作表的数据锚点，确定名称列、数据列和数据起始行位置。

    通过查找“读数1”定位读数标题行，计算数据起始行；通过“构件名称”关键词调整名称列，通过“读数1”调整数据列。

    Args:
        ws: Excel工作表对象（openpyxl.worksheet.worksheet.Worksheet）
    Returns:
        dict: 锚点信息字典，包含以下键：
            - name_col: 名称列索引（int）
            - data_col: 数据列起始索引（int）
            - data_row: 数据起始行索引（int）
            - read_row: 读数标题行索引（int）
    """
    read_row = None
    for r in range(1, 60):
        for c in range(1, 40):
            if "读数1" in str(ws.cell(row=r, column=c).value or ""):
                read_row = r;
                break
        if read_row: break
    data_start_row = (read_row + 1) if read_row else 7
    name_col = 2
    for r in range(1, (read_row or 15) + 1):
        for c in range(1, 30):
            if "构件名称" in str(ws.cell(row=r, column=c).value or ""):
                name_col = c;
                break
        if name_col != 2: break
    data_col = None
    if read_row:
        for c in range(1, 40):
            if "读数1" in str(ws.cell(row=read_row, column=c).value or ""):
                data_col = c;
                break
    data_col = data_col or 5
    return {"name_col": name_col, "data_col": data_col, "data_row": data_start_row, "read_row": read_row or 6}


def keep_align(cell, value):
    """
    向Excel单元格写入值并保留原有对齐格式，避免格式错乱。

    读取单元格原有对齐方式（水平/垂直对齐、自动换行等），写入值后重新应用这些格式。

    Args:
        cell: Excel单元格对象（openpyxl.cell.cell.Cell）
        value: 待写入的值（str）
    """
    old = cell.alignment or Alignment()
    cell.value = value
    cell.alignment = Alignment(
        horizontal=old.horizontal,
        vertical=old.vertical,
        wrap_text=old.wrap_text,
        textRotation=old.textRotation,
        indent=old.indent,
        shrinkToFit=old.shrinkToFit
    )


def write_block(ws, anchors, pos, item):
    """
    将数据块写入Excel工作表的指定位置，保留格式对齐。

    根据锚点信息计算起始行，写入构件名称和5行数据，确保与模板格式一致。

    Args:
        ws: Excel工作表对象（openpyxl.worksheet.worksheet.Worksheet）
        anchors: 锚点信息字典（detect_anchors返回结果）
        pos: 数据块在工作表中的位置（0-4，int）
        item: 数据块对象（expand_blocks返回的单个元素）
    """
    r0 = anchors["data_row"] + pos * PER_LINE_PER_BLOCK
    name_col = anchors["name_col"];
    data_col = anchors["data_col"]
    keep_align(ws.cell(row=r0, column=name_col), item["name"])
    for dr in range(PER_LINE_PER_BLOCK):
        for dc in range(9):
            ws.cell(row=r0 + dr, column=data_col + dc).value = item["data"][dr][dc]


def slash_block(ws, anchors, pos):
    """
    用“/”填充Excel工作表中指定位置的数据块，用于补齐未填满的区域。

    在指定位置写入“/”占位符，保留单元格原有对齐格式，确保表格格式统一。

    Args:
        ws: Excel工作表对象（openpyxl.worksheet.worksheet.Worksheet）
        anchors: 锚点信息字典（detect_anchors返回结果）
        pos: 数据块位置（0-4，int）
    """
    r0 = anchors["data_row"] + pos * PER_LINE_PER_BLOCK
    name_col = anchors["name_col"];
    data_col = anchors["data_col"]
    keep_align(ws.cell(row=r0, column=name_col), "/")
    for dr in range(PER_LINE_PER_BLOCK):
        for dc in range(9):
            ws.cell(row=r0 + dr, column=data_col + dc).value = "/"


def slash_tail(ws, anchors, used_pos):
    """
    用“/”填充工作表中未使用的数据块位置，从已用位置到最后。

    确保工作表数据区域格式统一，未使用的位置明确标记为“/”。

    Args:
        ws: Excel工作表对象（openpyxl.worksheet.worksheet.Worksheet）
        anchors: 锚点信息字典（detect_anchors返回结果）
        used_pos: 已使用的数据块位置索引（int）
    """
    for rem in range(used_pos, BLOCKS_PER_SHEET):
        slash_block(ws, anchors, rem)


# ===== 元信息固定坐标 / 仪器识别 =====
def top_left_of_merged(ws, r, c):
    """
    查找合并单元格的左上角单元格坐标，确保值写入正确位置。

    遍历工作表中的合并区域，返回指定单元格所属合并区域的左上角行号和列号。

    Args:
        ws: Excel工作表对象（openpyxl.worksheet.worksheet.Worksheet）
        r: 行号（int）
        c: 列号（int）
    Returns:
        tuple: 左上角单元格的行号和列号（int, int）
    """
    for rng in ws.merged_cells.ranges:
        if rng.min_row <= r <= rng.max_row and rng.min_col <= c <= rng.max_col:
            return rng.min_row, rng.min_col
    return r, c


def apply_meta_fixed(wb, categories_present, meta: dict):
    """
    向Excel工作表写入固定元信息（工程名称、委托编号）到指定位置。

    仅处理目标类型工作表，将工程名称写入C3、委托编号写入L3，支持合并单元格。

    Args:
        wb: Excel工作簿对象（openpyxl.workbook.Workbook）
        categories_present: 存在的构件类型列表（list[str]）
        meta: 元信息字典，含"proj"（工程名称）和"order"（委托编号）键
    """
    for ws in wb.worksheets:
        if not any(ws.title.startswith(p) for p in categories_present): continue

        def _set_rc(r, c, v):
            if not v: return
            r0, c0 = top_left_of_merged(ws, r, c)
            ws.cell(row=r0, column=c0).value = v

        _set_rc(3, 3, meta.get("proj"))  # C3
        _set_rc(3, 12, meta.get("order"))  # L3


def find_avg_col(ws, read_row_guess: int):
    """
    查找Excel工作表中平均值列的位置，优先匹配含“平均”和“厚”的列。

    在指定的读数标题行附近查找含“平均”和“厚”关键词的列，兜底返回M列（第13列）。

    Args:
        ws: Excel工作表对象（openpyxl.worksheet.worksheet.Worksheet）
        read_row_guess: 读数标题行猜测位置（int）
    Returns:
        int: 平均值列的列号（int）
    """
    for c in range(1, 50):
        v = str(ws.cell(row=read_row_guess, column=c).value or "")
        if "平均" in v and "厚" in v: return c
        if "平均值" in v: return c
    return 13  # 兜底 M 列


def detect_instrument(ws):
    """
    根据平均值列数据自动识别仪器型号（23-90或24-57）。

    规则：平均值<10 → 23-90；≥10 → 24-57，通过检查数据起始行后的前25行平均值判断。

    Args:
        ws: Excel工作表对象（openpyxl.worksheet.worksheet.Worksheet）
    Returns:
        str: 仪器型号，可能为"23-90"或"24-57"
    """
    anc = detect_anchors(ws)
    avg_col = find_avg_col(ws, anc["read_row"])
    start_r = anc["data_row"]
    end_r = min(start_r + 24, ws.max_row)
    for r in range(start_r, end_r + 1):
        v = ws.cell(row=r, column=avg_col).value
        if v is None: continue
        if isinstance(v, (int, float)):
            num = float(v);
            return "24-57" if num >= 10 else "23-90"
        s = str(v).strip()
        if s == "/": continue
        m = re.search(r"-?\d+(?:\.\d+)?", s)
        if m:
            num = float(m.group(0))
            return "24-57" if num >= 10 else "23-90"
    return "23-90"


def write_instrument(ws, text):
    """
    向Excel工作表写入仪器型号到固定位置（E33:H33合并区域）。

    定位E33:H33合并区域的左上角单元格，写入仪器型号。

    Args:
        ws: Excel工作表对象（openpyxl.worksheet.worksheet.Worksheet）
        text: 仪器型号字符串（str）
    """
    r0, c0 = top_left_of_merged(ws, 33, 5)  # E33:H33 合并左上
    ws.cell(row=r0, column=c0).value = text


def apply_meta_on_pages(wb, pages: list[str], date_str: str, env_str: str, auto_instrument=True):
    """
    向指定Excel工作表写入日期、环境温度和仪器型号元信息。

    日期写入K33，环境温度写入K34；若开启自动识别，仪器型号根据平均值列数据自动判断并写入。

    Args:
        wb: Excel工作簿对象（openpyxl.workbook.Workbook）
        pages: 工作表名称列表（list[str]）
        date_str: 日期字符串（str）
        env_str: 环境温度字符串（str）
        auto_instrument: 是否自动识别仪器型号，默认True
    """
    if not pages: return
    for name in pages:
        ws = wb[name]

        def _set_rc(r, c, v):
            if not v: return
            r0, c0 = top_left_of_merged(ws, r, c)
            ws.cell(row=r0, column=c0).value = v

        _set_rc(33, 11, date_str)  # K33
        _set_rc(34, 11, env_str)  # K34
        if auto_instrument:
            write_instrument(ws, detect_instrument(ws))


# ===== 规范化 =====
def normalize_date(text: str) -> str:
    """
    将用户输入的环境温度字符串规范化为“X℃”或“X.X℃”格式。

    从输入中提取数字部分（忽略“℃”“度”等符号），整数温度去小数点，小数温度保留有效数字。
    若无法提取有效数字，则返回原始字符串。

    Args:
        text: 用户输入的环境温度字符串（如“24”“24℃”“24.5度”）
    Returns:
        str: 标准化的温度字符串（如“24℃”“24.5℃”）
    """
    s = (text or "").strip()
    if not s: return ""
    if re.fullmatch(r"\d{8}", s):
        y, m, d = int(s[:4]), int(s[4:6]), int(s[6:8]);
        return f"{y}年{m}月{d}日"
    s2 = s.replace("年", " ").replace("月", " ").replace("日", " ")
    for ch in ".-/，,": s2 = s2.replace(ch, " ")
    nums = re.findall(r"\d+", s2)
    if len(nums) >= 3:
        y, m, d = map(int, nums[:3]);
        return f"{y}年{m}月{d}日"
    return s


def normalize_env(text: str) -> str:
    """
    将用户输入的环境温度字符串规范化为“X℃”或“X.X℃”格式。

    从输入中提取数字部分（忽略“℃”“度”等符号），整数温度去小数点，小数温度保留有效数字。
    若无法提取有效数字，则返回原始字符串。

    Args:
        text: 用户输入的环境温度字符串（如“24”“24℃”“24.5度”）
    Returns:
        str: 标准化的温度字符串（如“24℃”“24.5℃”）
    """
    s = (text or "").strip()
    if not s: return ""
    m = re.search(r"-?\d+(?:\.\d+)?", s)
    if not m: return s
    val = float(m.group(0))
    return f"{int(val)}℃" if val.is_integer() else f"{str(val).rstrip('0').rstrip('.')}℃"


def _normalize_date_token(tok: str, base_year: int) -> str:
    """将单个日期 token 规范为"YYYY-MM-DD"，失败返回空串。"""
    if not tok:
        return ""
    tok = tok.strip()
    tok = tok.replace("年", "-").replace("月", "-").replace("日", "")
    tok = tok.replace("/", "-").replace(".", "-")
    tok = re.sub(r"\s+", "-", tok)
    if re.fullmatch(r"\d{8}", tok):
        y = int(tok[:4]);
        mth = int(tok[4:6]);
        d = int(tok[6:])
    else:
        m = re.fullmatch(r"(\d{4})-(\d{1,2})-(\d{1,2})", tok)
        if m:
            y, mth, d = map(int, m.groups())
        else:
            m = re.fullmatch(r"(\d{1,2})-(\d{1,2})", tok)
            if not m:
                return ""
            y = base_year
            mth, d = map(int, m.groups())
    if not (1 <= mth <= 12 and 1 <= d <= 31):
        return ""
    return f"{y:04d}-{mth:02d}-{d:02d}"


def _parse_dates_simple(input_str: str):
    """简单解析多个日期，返回 (日期列表, 无效token列表)。"""
    # 允许空格/英文逗号/中文逗号/中文顿号作为分隔
    tokens = [t for t in re.split(r"[,\s，、]+", input_str.strip()) if t]

    res, ignored = [], []
    seen = set()
    base_year = None
    cur_year = datetime.now().year

    i = 0
    while i < len(tokens):
        tok = tokens[i]
        consumed = 1

        # 先尝试把当前 token 当成一个完整日期（支持 8/27、8-27、2025-8-27、2025年8月27日 等等）
        norm = _normalize_date_token(tok, base_year or cur_year)

        if not norm:
            # 尝试 Y M D 这种被空格/逗号拆开的情况：2025 8 27
            if re.fullmatch(r"\d{4}", tok) and i + 2 < len(tokens) \
                    and tokens[i + 1].isdigit() and tokens[i + 2].isdigit():
                norm = _normalize_date_token(
                    f"{tok}-{tokens[i + 1]}-{tokens[i + 2]}",
                    base_year or cur_year
                )
                consumed = 3

            # 尝试 M D：8 27（基于 base_year 或当前年）
            elif tok.isdigit() and i + 1 < len(tokens) and tokens[i + 1].isdigit():
                norm = _normalize_date_token(
                    f"{tok}-{tokens[i + 1]}",
                    base_year or cur_year
                )
                consumed = 2

        if norm:
            # 锁定 base_year，后续 M-D 走同一年
            if base_year is None:
                base_year = int(norm[:4])
            # 去重：同一天不重复计入
            if norm not in seen:
                res.append(norm)
                seen.add(norm)
        else:
            # 记录无法解析的原始 token（或组合）
            ignored.extend(tokens[i:i + consumed])

        i += consumed

    return res, ignored

    # ===== 交互 =====


HELP_HOME = f"""
====================  The Unification | 帮助中心  ====================
this application was made by {AUTHOR} in 2025 summer
使用方式：
  • 在“请输入 Word 源路径”处，输入 help 打开本帮助中心
  • 在本界面输入 1 / 2 / 3 / 4 查看对应模式的完整教程
  • 直接回车 返回到路径输入界面
  • 在任何步骤输入小写 q 可返回上一步；仅在路径输入界面输入大写 Q 退出程序

全局规则（适用于所有模式）：
  • 日期输入支持以下格式（可混用，自动标准化）：
      YYYY-MM-DD / YYYY/MM/DD / YYYY.MM.DD / YYYY MM DD / YYYYMMDD
      M-D / M/D / M.D / M D / YYYY年M月D日
  • 温度输入：任意字符串（如 24℃ / 24.5 度），自动标准化为“X℃”或“X.X℃”
  • “支撑（WZ）”分桶策略（仅 Mode 1/2/3）：
      - 在进入“支撑”配置之前询问：1=按编号；2=按楼层（与钢柱/钢梁一致）
  • 输出规则：
      - 统一使用模板页池命名（不在 Sheet 名称中写日期/楼层）
      - 日期写入 K33、温度写入 K34，仪器型号自动识别写入（E33:H33）
  • 排序规则：
      - 楼层自然顺序：B* → 1F↑ → 机房层 → 屋面
      - 同层内：WZ编号（若有）→ 名称里的数字 → 名称字典序（稳定、可复现）

提示：
  • 任何模式完成或发生错误后，程序都会回到路径输入界面
  • 仅当在路径输入界面输入大写 Q，程序才会退出
=====================================================================
"""

HELP_TEXTS = {
    "1":
        """====================  Mode 1 | 按日期分桶（默认稳健）  ====================
        
        适用场景：
          将全部构件按日期分配到多天；支持“后面的日子优先”或“前面的日子优先”。
        
        操作流程：
          1) 选择模式：输入 1
          2) 若存在“支撑”，在进入“支撑”配置之前选择分桶策略：
               - 1 = 按编号（WZ号）   2 = 按楼层（与钢柱/钢梁一致）
          3) 录入“日期桶”（1~10 天；日期格式见帮助首页）
          4) 选择规则重叠优先级：
               - 回车 = “后面的日子优先”（默认），n = “前面的日子优先”
          5) 预览分配结果：
               - 回车 = 确认生成
               - n    = 取消
               - a    = 将未分配构件并入最后一天
          6) 系统按天写入页池并批量写元信息（日期、温度、仪器）
        
        输出与命名：
          • 工作表命名沿用模板页池（“钢柱/钢梁/支撑/其他（n）”）
          • 日期写 K33，温度写 K34；仪器型号自动识别
        
        输入示例：
          2025-08-27, 2025/8/28, 20250829
        
        返回/退出：
          • 任意步骤输入 q 返回上一步
          • 完成或出错后，自动回到路径输入；仅在路径输入处输入 Q 才退出
        =====================================================================
        """,
    "2":
        """====================  Mode 2 | 按楼层断点（按层出报）  ====================
        
        适用场景：
          按“楼层范围”定义若干桶（如 1F-3F、4F-6F、B3-B1、屋面/机房层），
          每个桶映射到一天（或多天）与温度。
        
        操作流程：
          1) 选择模式：输入 2
          2) 若存在“支撑”，在进入“支撑”配置之前选择分桶策略：
               - 1 = 按编号（WZ号）   2 = 按楼层（与钢柱/钢梁一致）
          3) 定义楼层桶（顺序自动规范：B* → 1F↑ → 机房层 → 屋面）
          4) 为每个桶指定日期与（可选）温度；也可按需使用统一设置
          5) 预览 → 确认 → 写入
        
        输出与命名：
          • 工作表命名沿用模板页池；日期/温度写入 K33/K34；仪器自动识别
        
        输入示例：
          桶：B3-B1、1F-5F、屋面
          日期：2025.8.29
        
        返回/退出：
          • 任意步骤输入 q 返回上一步
          • 完成或出错后，自动回到路径输入；仅在路径输入处输入 Q 才退出
        =====================================================================
        """,
    "3":
        """====================  Mode 3 | 单日模式（最简方案）  ====================
        
        适用场景：
          全量构件归入同一日期与温度；快速制表或整单同日检测。
        
        操作流程：
          1) 选择模式：输入 3
          2) 若存在“支撑”，在进入“支撑”配置之前选择分桶策略：
               - 1 = 按编号（WZ号）   2 = 按楼层（与钢柱/钢梁一致）
          3) 输入日期与（可选）温度
          4) 写入页池；自动分页（25 行/页 = 5 组 × 5 行）
        
        输出与命名：
          • 工作表命名沿用模板页池；日期/温度写入 K33/K34；仪器自动识别
        
        输入示例：
          20250101   2025年1月1日   2025 1 1
        
        返回/退出：
          • 任意步骤输入 q 返回上一步
          • 完成或出错后，自动回到路径输入；仅在路径输入处输入 Q 才退出
        =====================================================================
        """,
    "4":
        """================  Mode 4 | 楼层 × 日期 切片（灵活均分/配额）  ================
        
        适用场景：
          同一楼层需要分配到多天；可选择“均分”或“每日上限（配额）”进行切片。
        
        核心概念：
          • 共用计划：为一批选定楼层设置“同一套”日期清单与每日上限（空=均分）
          • 默认计划（*）：为“未单独配置”的楼层设置的通用计划
          • 兜底：若仍有未分配数据，可二选一：
              A) 统一日期/温度一次性分配；或
              B) 回落到 Mode 1（日期分桶）流程
        
        操作流程：
          1) 选择模式：输入 4
          2) 选择适用楼层（留空=全部已识别楼层；支持 B2、5F、屋面、机房层）
          3) 是否“共用计划”：
              - y = 共用：一次录入日期清单与每日上限（空=均分），套用到所有选定楼层
              - 回车 = 分别设置：按楼层逐一录入日期与每日上限
          4) 存在“未配置楼层”时，是否创建默认计划（*）：
              - y = 创建：再录入一次日期与每日上限，通用于剩余楼层
              - 回车 = 不创建：留待后续兜底
          5) 分发与兜底：
              - 已配置楼层：立即切片、分页、写入
              - 未配置楼层：选择统一日期/温度一次性分配，或回落到 Mode 1 分桶流程
        
        输出与命名：
          • 工作表命名沿用模板页池；按“日期切片”分组批量写入 K33/K34；仪器自动识别
          • 顺序稳定：楼层排序 B* → 1F↑ → 机房层 → 屋面；同层内按 WZ 编号 → 数字 → 字典序
        
        输入示例：
          共用：楼层 5F, 6F, B2；日期 2025-08-27, 20250828, 2025年8月29日；上限 60
          分别：5F → 8/27, 8/28（上限空=均分）；6F → 2025.8.27（上限 40）
        
        返回/退出：
          • 任意步骤输入 q 返回上一步
          • 完成或出错后，自动回到路径输入；仅在路径输入处输入 Q 才退出
        =====================================================================
        """,
}


def tutorial_browser():
    """显示模式教程浏览器。"""
    print(HELP_HOME)
    viewed = False
    while True:
        prompt = "还要查看其他模式？输入 1/2/3/4，回车或 q 返回。\n→ " if viewed else "查看哪个模式？输入 1/2/3/4，回车或 q 返回路径输入。\n→ "
        sel = input(prompt).strip()
        if sel in ("", "q"):
            return
        if sel in HELP_TEXTS:
            print(HELP_TEXTS[sel])
            viewed = True
        else:
            print("仅接受 1/2/3/4 或回车/q。")


def prompt_path(prompt, default: Path) -> Path:
    """
    交互式获取用户输入文件路径，验证文件存在性并返回有效路径。

    提示用户输入文件路径，支持直接回车使用默认路径；自动处理路径中的引号；
    若输入路径无效（文件不存在），则显示错误提示并重新请求输入，确保返回有效文件路径。

    Args:
        prompt: 路径输入提示信息（str）
        default: 默认文件路径（Path对象）
    Returns:
        Path: 经过验证的有效文件路径
    """
    while True:
        raw = ask(f"{prompt}（回车默认：{default}）")
        if raw.lower() == "help":
            tutorial_browser()
            continue
        p = Path(raw.strip('"')) if raw else default
        if p.exists() and p.is_file():
            return p
        print(f"❌ 找不到文件：{p}")


def prompt_floor_breaks(label: str):
    """
    交互式获取楼层断点列表，支持无效输入并返回空值处理。

    提示用户输入空格分隔的楼层断点（如"5 10"），支持直接回车表示不分段；
    自动过滤重复值并按升序排序；若输入无效（非数字）则返回空列表。

    Args:
        label: 提示信息前缀（str）
    Returns:
        list[int]: 排序后的楼层断点列表（空列表表示不分段）
    """
    txt = ask(f"{label} 断点楼层（空格分隔，如 5 10；回车=不分段）：")
    if not txt: return []
    try:
        return sorted({int(x) for x in txt.split()})
    except:
        return []


# ===== 日期分桶（泛化到任意类别） =====
def _parse_int_ranges(rule: str):
    """
    解析整数范围表达式为整数列表，支持单值和范围格式。

    支持的格式示例：
    - 单值："5" → [5]
    - 范围："3-7" → [3,4,5,6,7]
    - 混合："2,5-7,9" → [2,5,6,7,9]

    Args:
        rule: 包含整数或范围的字符串（如"3-7,9"）
    Returns:
        list[int]: 解析后的整数列表（按升序排列）
    """
    res = []
    if not rule.strip(): return res
    for tok in re.split(r"[,\s，]+", rule.strip()):
        if not tok: continue
        m = re.match(r"^\s*(\d+)\s*-\s*(\d+)\s*$", tok)
        if m:
            a, b = int(m.group(1)), int(m.group(2))
            if a > b: a, b = b, a
            res.append((a, b))
        else:
            m = re.match(r"^\d+$", tok)
            if m:
                v = int(tok);
                res.append((v, v))
            else:
                m = re.match(r"(?i)^[FL]\s*(\d+)$", tok) or re.match(r"(?i)^(\d+)\s*[FL]$", tok)
                if m:
                    v = int(m.group(1));
                    res.append((v, v))
                else:
                    lt = tok.lower()
                    if lt.startswith(("屋面", "屋顶层", "顶层", "wm", "dc")):
                        res.append((10 ** 6, 10 ** 6))
                    elif lt.startswith(("机房层", "jf")):
                        res.append((10 ** 6 - 1, 10 ** 6 - 1))
    return res


def parse_rule(text: str):
    """
    解析数据分发规则字符串为结构化规则字典。

    支持两种规则类型：
    - 启用所有数据：输入“*”“all”“全部”“所有”时，返回启用状态且空范围（表示接收所有数据）
    - 范围规则：其他输入解析为整数范围列表（通过_parse_int_ranges处理）

    Args:
        text: 规则字符串（如“*”“1-3 5”“全部”）
    Returns:
        dict: 规则字典，包含：
            - enabled: 是否启用该规则（bool）
            - ranges: 解析后的范围列表（list[tuple[int, int]]，空列表表示全部）
    """
    s = (text or "").strip()
    if not s: return {"enabled": False, "ranges": []}
    if s.lower() in ("*", "all") or s in ("全部", "所有"): return {"enabled": True, "ranges": []}
    return {"enabled": True, "ranges": _parse_int_ranges(s)}


def _in_ranges(val: int, ranges):
    """
    判断值是否在指定的范围列表内，支持空范围表示“全部包含”。

    范围列表为空时默认包含所有值；否则检查值是否落在任一范围的闭区间内。

    Args:
        val: 待判断的整数（如楼层号、支撑编号）
        ranges: 范围元组列表（如[(1,3), (5,7)]），空列表表示全部
    Returns:
        bool: 在范围内返回True，否则返回False
    """
    if ranges is None: return False
    if ranges == []: return True  # noqa
    for a, b in ranges:
        if a <= val <= b: return True
    return False


def _wz_no(name: str):
    """
    从支撑构件名称中提取编号（如从“WZ3”“支撑-5”中提取3、5）。

    支持关键词匹配：
    - 含“WZ”或“ZC”前缀（如“WZ12”“ZC-8”）
    - 含“支撑”关键词（如“支撑6”“斜撑-3”）
    提取失败时返回None。

    Args:
        name: 支撑构件名称字符串（如“WZ5”“支撑-10”）
    Returns:
        int | None: 提取的编号，失败则返回None
    """
    m = re.search(r"(?i)\b(?:WZ|ZC)\s*[-–—]?\s*(\d+)\b", name)
    if m: return int(m.group(1))
    m = re.search(r"支撑\s*[-–—]?\s*(\d+)", name)
    return int(m.group(1)) if m else None


def _match_keywords(name: str, kws):
    """
    判断构件名称是否包含任意关键词（忽略大小写）。

    关键词列表为空时默认匹配所有名称；否则检查名称是否含任一关键词（不区分大小写）。

    Args:
        name: 构件名称字符串
        kws: 关键词列表（如["3层", "东立面"]）
    Returns:
        bool: 包含任一关键词返回True，否则返回False（关键词为空时返回True）
    """
    if not kws: return True
    s = name.lower()
    return any(k.lower() in s for k in kws)


def prompt_mode():
    """模式选择，支持 q 返回。"""
    txt = ask("模式选择：1) 按日期分桶  2) 按楼层断点  3) 单日模式  4) 楼层+日期配额")
    if txt in ("", "1"):
        return "1"
    if txt in ("2", "3", "4"):
        return txt
    return "1"


def prompt_bucket_priority():
    """询问规则重叠优先级。"""
    ans = ask("规则重叠将按【后面的天】优先并自动做差（回车=是 / n=否）：", lower=True)
    return ans != 'n'


def prompt_support_strategy_for_bucket():
    """在需要支撑分桶策略时询问一次。"""
    global support_bucket_strategy
    if support_bucket_strategy is None:
        ans = ask("支撑分桶方式：1) 按编号 2) 按楼层（回车=1）")
        support_bucket_strategy = "floor" if ans == "2" else "number"
    return support_bucket_strategy


def prompt_date_buckets(categories_present):
    """
    交互式收集日期桶配置，支持1-10天的检测数据分发规则。

    为每天配置：
    - 日期（自动标准化为“YYYY年MM月DD日”）
    - 环境温度（自动标准化为“X℃”）
    - 各构件类型的接收规则（楼层/编号范围）
    - 关键词筛选（可选）

    Args:
        categories_present: 存在的构件类型列表（如["钢柱", "支撑"]）
    Returns:
        list[dict]: 日期桶配置列表，每个元素含日期、环境、规则等信息
    """
    while True:
        n_txt = ask("共有几天（1-10，回车=1）：")
        if not n_txt: n = 1; break
        if n_txt.isdigit() and 1 <= int(n_txt) <= 10:
            n = int(n_txt);
            break
        print("请输入 1-10 之间的整数。")
    buckets = []
    for i in range(1, n + 1):
        print(f"\n—— 第 {i} 天 ——")
        d = ask("📅 日期（20250101 / 2025年1月1日 / 2025 1 1 / 2025.1.1 / 2025-1-1 / 1-1 / 01-01）：")
        e = ask("🌡 环境温度（24 / 24℃ / 24 度 / 24 C）：")
        rules = {}
        for cat in categories_present:
            if cat == "支撑":
                prompt_support_strategy_for_bucket()
                if support_bucket_strategy == "floor":
                    txt = ask("🦾 支撑 楼层规则（例：1-3 5 7-10 屋面；留空=不接收；*=不限）：")
                else:
                    txt = ask("🦾 支撑 编号范围（例：1-12 20-25；留空=不接收；*=不限）：")
                    rules[cat] = parse_rule(txt)
            else:
                txt = ask(f"🏗 {cat} 楼层规则（例：1-3 5 7-10 屋面；留空=不接收；*=不限）：")
                rules[cat] = parse_rule(txt)
        kws_txt = ask("🔎 关键词（可多个，空格/逗号分隔；留空=无需）：")
        buckets.append({
            "date_raw": d,
            "date": normalize_date(d) if d else "",
            "env": normalize_env(e) if e else "",
            "rules": rules,
            "kws": [k for k in re.split(r"[,\s，]+", kws_txt) if k] if kws_txt else []
        })
    return buckets


def assign_by_buckets(cat_groups: dict, buckets, later_priority=True):
    """
    将构件数据组按日期桶规则分配到对应天数，支持规则重叠处理。

    分配逻辑：
    1. 按构件类型遍历数据组
    2. 根据日期桶规则（楼层/编号范围+关键词）匹配数据
    3. 规则重叠时按“后定义桶优先”（可通过参数关闭）
    返回分配结果和未匹配的数据。

    Args:
        cat_groups: 按类型分组的构件数据（键为类型，值为数据组列表）
        buckets: 日期桶配置列表（prompt_date_buckets返回结果）
        later_priority: 规则重叠时是否后定义桶优先，默认True
    Returns:
        tuple: 包含两个元素的元组：
            - cat_byb: 按类型和桶分配的结果（dict[类型][桶索引] = 数据组列表）
            - remain_by_cat: 未分配的数据（dict[类型] = 数据组列表）
    """
    # 输出：cat_byb[cat][bucket_index] = [groups...];  remain_by_cat[cat] = [groups...]
    cat_byb = {cat: {i: [] for i in range(len(buckets))} for cat in cat_groups}
    assigned = {cat: set() for cat in cat_groups}
    order = range(len(buckets) - 1, -1, -1) if later_priority else range(len(buckets))
    for cat, groups in cat_groups.items():
        for idx, g in enumerate(groups):
            # 计算匹配
            fl = floor_of(g["name"])
            wzno = _wz_no(g["name"]) if cat == "支撑" and support_bucket_strategy == "number" else None
            for bi in order:
                b = buckets[bi]
                rule = b["rules"].get(cat, {"enabled": False, "ranges": None})
                if not rule.get("enabled"):
                    continue
                ok = False  # noqa
                if cat == "支撑":
                    if support_bucket_strategy == "number":
                        rng = rule["ranges"]
                        ok_num = True if rng == [] else (wzno is not None and _in_ranges(wzno, rng))
                        ok = ok_num
                    else:
                        ok = _in_ranges(fl, rule["ranges"])
                else:
                    ok = _in_ranges(fl, rule["ranges"])
                if ok and _match_keywords(g["name"], b["kws"]):
                    cat_byb[cat][bi].append(g);
                    assigned[cat].add(idx);
                    break

    remain_by_cat = {cat: [g for i, g in enumerate(groups) if i not in assigned[cat]]
                     for cat, groups in cat_groups.items()}
    return cat_byb, remain_by_cat


def preview_buckets_generic(cat_byb, remain_by_cat, buckets, categories_present):
    """
     预览日期桶分配结果，询问用户是否确认生成，支持未分配数据处理。

     显示每天各类型构件的分配数量及未分配数据；提供选项：
     - 回车：确认生成
     - n：取消操作
     - a：将未分配数据并入最后一天

     Args:
         cat_byb: 按类型和桶分配的结果
         remain_by_cat: 未分配数据
         buckets: 日期桶配置列表
         categories_present: 存在的构件类型列表
     Returns:
         tuple: 包含两个元素的元组：
             - 是否确认生成（bool）
             - 是否将未分配数据并入最后一天（bool）
     """
    print("\n🧾 预览：")
    for i, b in enumerate(buckets, start=1):
        parts = []
        for cat in categories_present:
            parts.append(f"{cat} {len(cat_byb[cat][i - 1])}")
        print(f"  第{i}天 〔{b['date'] or b['date_raw'] or '未填日期'} / {b['env'] or '未填温度'}〕 → " + "、".join(parts))
    if any(remain_by_cat[cat] for cat in categories_present):
        print("  ⚠️ 未分配：", end="")
        print("、".join(f"{cat} {len(remain_by_cat[cat])}" for cat in categories_present if remain_by_cat[cat]))
    ans = ask("确认生成吗？(回车=是 / n=否 / a=把未分配并入最后一天)：", lower=True)
    return (ans != "n"), (ans == "a")


def expand_blocks_by_bucket(cat_byb):
    """
    将按日期桶分配的构件数据组拆分为标准数据块（5行/块）。

    对每个类型、每个日期桶的数据组应用expand_blocks函数，确保数据块结构统一，适配Excel模板。

    Args:
        cat_byb: 按类型和桶分配的结果（assign_by_buckets返回的cat_byb）
    Returns:
        dict: 按类型和桶组织的数据块字典（dict[类型][桶索引] = 数据块列表）
    """
    # 返回：blocks_by_cat[cat][bucket_index] = [blocks...]
    return {cat: {bi: expand_blocks(lst, PER_LINE_PER_BLOCK) for bi, lst in byb.items()}
            for cat, byb in cat_byb.items()}


def ensure_pages_slices_for_cat(wb, cat: str, blocks_by_bucket_for_cat: dict):
    """
    为指定类型的每个日期桶确保足够的工作表，返回按桶划分的工作表切片。

    计算每个桶所需工作表数量（按5块/页），不足时自动复制补充：
    - 常规类型（钢柱/钢梁/支撑）从自身基础表复制
    - “其他”类型从钢柱模板复制
    返回按桶分组的工作表名称列表。

    Args:
        wb: Excel工作簿对象
        cat: 构件类型（如“钢柱”“其他”）
        blocks_by_bucket_for_cat: 该类型按桶组织的数据块字典
    Returns:
        list[list[str]]: 按桶划分的工作表名称列表（每个元素为一个桶的工作表）
    """

    def need_pages(lst):
        return math.ceil(len(lst) / BLOCKS_PER_SHEET) if lst else 0

    page_need_each = [need_pages(blocks_by_bucket_for_cat.get(i, [])) for i in range(len(blocks_by_bucket_for_cat))]
    total_need = sum(page_need_each)
    if total_need == 0:
        return [[] for _ in page_need_each]
    if cat == "其他":
        pages_all = ensure_total_pages_from(wb, "钢柱", "其他", total_need)
    else:
        pages_all = ensure_total_pages(wb, cat, total_need)
    slices = [];
    p = 0
    for n in page_need_each:
        slices.append(pages_all[p:p + n]);
        p += n
    return slices


def make_target_order_generic(pages_slices_by_cat, categories_present):
    """
    生成工作表的目标顺序，按“日期桶→类型优先级”排序。

    排序规则：
    1. 按日期桶轮次分组
    2. 同轮次内按CATEGORY_ORDER（钢柱→钢梁→支撑→其他）排序
    确保工作表按检测流程和类型逻辑有序排列。

    Args:
        pages_slices_by_cat: 按类型和桶划分的工作表切片字典
        categories_present: 存在的构件类型列表
    Returns:
        list[str]: 排序后的工作表名称列表
    """
    rounds = 0
    for cat in categories_present:
        rounds = max(rounds, len(pages_slices_by_cat.get(cat, [])))
    target = []
    for i in range(rounds):
        for cat in CATEGORY_ORDER:
            if cat not in categories_present: continue
            sl = pages_slices_by_cat[cat][i] if i < len(pages_slices_by_cat[cat]) else []
            target += sl
    return target


# ===== Excel 写入带进度 =====
class Prog:
    def __init__(self, total: int, label: str = "写入 Excel"):
        self.total = max(1, total)
        self.done = 0
        self.label = label

    def tick(self, k=1):
        self.done += k
        pct = int(self.done * 100 / self.total)
        sys.stdout.write(f"\r📊 {self.label}：{self.done}/{self.total}（{pct}%）")
        sys.stdout.flush()

    def finish(self):
        sys.stdout.write("\n");
        sys.stdout.flush()


def fill_blocks_to_pages(wb, pages_slice, blocks, prog: Prog | None = None):
    """
    将数据块填充到指定的Excel工作表，支持进度跟踪。

    按工作表顺序填充数据块，每页最多5个块；页面填满后自动切换到下一页；
    未填满的页面用“/”补齐空白区域；支持通过Prog对象跟踪进度。

    Args:
        wb: Excel工作簿对象
        pages_slice: 工作表名称列表（当前桶的工作表）
        blocks: 待填充的数据块列表
        prog: 进度跟踪对象（可选）
    """
    if not pages_slice: return
    page_idx, pos = 0, 0
    for it in blocks:
        if page_idx >= len(pages_slice): break
        ws = wb[pages_slice[page_idx]]
        anc = detect_anchors(ws)
        write_block(ws, anc, pos, it)
        if prog: prog.tick(1)
        pos += 1
        if pos == BLOCKS_PER_SHEET:
            page_idx += 1;
            pos = 0
    if page_idx < len(pages_slice) and pos != 0:
        ws = wb[pages_slice[page_idx]]
        slash_tail(ws, detect_anchors(ws), pos)


def cleanup_unused_sheets(wb, used_names, bases=("钢柱", "钢梁", "支撑", "其他")):
    """
    清理Excel中未使用的指定类型工作表，减少冗余。

    仅保留已使用的目标类型工作表（钢柱/钢梁/支撑/其他），避免模板中多余工作表干扰。
    确保至少保留一个工作表（防止工作簿为空）。

    Args:
        wb: Excel工作簿对象
        used_names: 已使用的工作表名称列表
        bases: 目标类型基础名称列表
    """
    used = set(used_names)
    to_remove = []
    for ws in list(wb.worksheets):
        if any(ws.title == b or ws.title.startswith(f"{b}（") for b in bases):
            if ws.title not in used:
                to_remove.append(ws)
    if len(to_remove) >= len(wb.worksheets):
        to_remove = to_remove[:-1]
    for ws in to_remove:
        wb.remove(ws)


def _distribute_by_dates(items, date_entries):
    """按日期列表将项目分配到各天。"""
    res = []
    if not date_entries:
        return res
    if date_entries[0][1] is not None:  # 配额模式
        cursor = 0
        total = len(items)
        for i, (d, limit, env) in enumerate(date_entries):
            if i < len(date_entries) - 1:
                take = min(limit, total - cursor)
            else:
                take = total - cursor
            res.append((d, env, items[cursor:cursor + take]))
            cursor += take
    else:  # 均分
        days = len(date_entries)
        per = math.ceil(len(items) / days) if days else 0
        cursor = 0
        for i, (d, _, env) in enumerate(date_entries):
            if i < days - 1:
                take = min(per, len(items) - cursor)
            else:
                take = len(items) - cursor
            res.append((d, env, items[cursor:cursor + take]))
            cursor += take
    return res


def _prompt_dates_and_limits():
    """交互获取日期、每日数量及环境温度。"""
    while True:
        txt = ask(
            "日期（空格/逗号分隔；支持 20250101 / 2025年1月1日 / 2025 1 1 / 2025.1.1 / 2025-1-1 / 1-1 / 01-01，\n"
            "年份默认取首个日期的年或当前年）：例如 2025-08-27 8-28 2025年1月1日\n→ "
        )
        if any(ch in txt for ch in "；;，、/\\|"):
            print("只接受逗号或空格分隔，请重输。")
            continue
        dates, ig = _parse_dates_simple(txt)
        if not dates:
            print("请输入至少一个合法日期。")
            continue
        if ig:
            print("已忽略：" + "、".join(ig))
        break
    while True:
        txt = ask("每日数量（按日期顺序；空=均分；填整数=配额）\n→ ")
        if txt == "":
            limits = [None] * len(dates)
            break
        tokens = [t for t in re.split(r"[ ,]+", txt) if t]
        if all(t.isdigit() and int(t) > 0 for t in tokens):
            if len(tokens) == 1:
                limits = [int(tokens[0])] * len(dates)
                break
            if len(tokens) == len(dates):
                limits = [int(t) for t in tokens]
                break
        print(f"请输入{len(dates)}个正整数或留空。")
    envs = []
    for d in dates:
        envs.append(ask(f"{d} 的环境温度（回车=不写）：\n→ "))
    return list(zip(dates, limits, envs))


def _summarize_plan(tag, plan, all_floors=None):
    """输出楼层计划摘要，便于用户确认。"""

    def fmt(entry):
        ds = " ".join(normalize_date(x[0]) for x in entry)
        ls = ",".join(str(x[1]) if x[1] is not None else "-" for x in entry)
        return f"{ds} → {ls}"

    specified = [f for f in plan if f != "*"]
    if specified:
        print("已单独配置：")
        for f in sorted(specified, key=_floor_sort_key_by_label):
            print(f"  {f} → {fmt(plan[f])}")
    if "*" in plan:
        print("默认配置：")
        print(f"  * → {fmt(plan['*'])}")
    if all_floors:
        miss = [f for f in all_floors if f not in plan and "*" not in plan]
        if miss:
            miss_txt = " ".join(sorted(miss, key=_floor_sort_key_by_label))
            print(f"未覆盖的楼层：{miss_txt} （稍后统一处理/回落到日期分桶）")


def _prompt_plan_for_floors(floors, shared=True):
    """针对给定楼层集合交互生成计划。"""
    floors = sorted(set(floors), key=_floor_sort_key_by_label)
    if floors:
        print("已识别楼层：" + " ".join(floors))
    # Step1 楼层
    while True:
        txt = ask("适用楼层（回车=全部）：示例 5F, 6F, B2, 屋面 或 5 6 B2\n→ ")
        if any(ch in txt for ch in "；;，、/\\|"):
            print("只接受逗号或空格分隔，请重输。")
            continue
        if not txt:
            sel = None
            break
        tokens = [t for t in re.split(r"[ ,]+", txt) if t]
        seen, sel, ig = set(), [], []
        for t in tokens:
            lb = _floor_label_from_name(t)
            if lb != "F?" and lb in floors and lb not in seen:
                sel.append(lb);
                seen.add(lb)
            else:
                ig.append(t)
        if ig:
            print("已忽略：" + "、".join(ig))
        if sel:
            break
        print("没有合法楼层，请重输。")
    targets = floors if sel is None else sel
    if shared:
        print("下面输入的日期与每日上限，将自动应用到以上所有楼层")
        date_entries = _prompt_dates_and_limits()
        if sel is None:
            return {"*": date_entries}
        return {f: date_entries for f in targets}
    plan = {}
    for f in targets:
        print(f"{f}：")
        plan[f] = _prompt_dates_and_limits()
    return plan


def prompt_mode4_plan(floors_by_cat, categories_present):
    """模式4交互，分别为各类别获取楼层计划。"""
    print("各类别楼层：")
    for cat in categories_present:
        fls = sorted(floors_by_cat.get(cat, set()), key=_floor_sort_key_by_label)
        print(f"{cat}: {(' '.join(fls)) if fls else '/'}")
    plans = {}
    for cat in categories_present:
        fls = floors_by_cat.get(cat, set())
        if not fls:
            continue
        print(f"\n[{cat}]")
        share = ask("这些楼层用同一套日期/数量吗？（y=是，回车=分别设置）\n→ ") == "y"
        plans[cat] = _prompt_plan_for_floors(fls, shared=share)
        # —— 新增：给未指定楼层兜底 ——
        all_floors = sorted(floors_by_cat.get(cat, set()), key=_floor_sort_key_by_label)
        plan_for_cat = plans[cat]
        specified = {f for f in plan_for_cat.keys() if f != "*"}
        if "*" not in plan_for_cat and len(specified) < len(all_floors):
            miss = [f for f in all_floors if f not in specified]
            print(f"👉 {cat} 还有未配置楼层：{' '.join(miss)}")
            ans = ask(
                "要不要给“未配置”的楼层用一套通用的日期/数量？（y=是，回车=跳过；未配置的楼层稍后会再统一询问或回落到日期分桶）",
                lower=True
            )
            if ans == "y":
                plan_for_cat["*"] = _prompt_dates_and_limits()
        _summarize_plan(cat, plan_for_cat, all_floors)
    return plans


def mode4_run(wb, grouped, categories_present):
    """执行模式4：按楼层和日期写入Excel。"""
    cf_groups = defaultdict(list)
    floors_by_cat = defaultdict(set)
    for cat in categories_present:
        for g in grouped[cat]:
            fl = _floor_label_from_name(g["name"])
            cf_groups[(cat, fl)].append(g)
            floors_by_cat[cat].add(fl)
    plan_dict = prompt_mode4_plan(floors_by_cat, categories_present)

    blocks_by_cat_bucket = {cat: defaultdict(list) for cat in CATEGORY_ORDER}
    buckets = []  # list[{date, env}]
    date_idx = {}
    env_by_date = {}
    leftover_by_cat = defaultdict(list)

    for (cat, fl), items in cf_groups.items():
        items.sort(key=lambda x: (
            int(re.search(r"\d+", x["name"]).group()) if re.search(r"\d+", x["name"]) else 10 ** 9, x["name"]))
        plan_for_cat = plan_dict.get(cat, {})
        plan = plan_for_cat.get(fl) or plan_for_cat.get("*")
        if not plan:
            leftover_by_cat[cat].extend(items)
            continue
        for date, env, slice_items in _distribute_by_dates(items, plan):
            if not slice_items:
                continue
            if date not in date_idx:
                date_idx[date] = len(buckets)
                buckets.append({"date": date, "env": env})
                env_by_date[date] = env
            elif env_by_date[date] != env:
                print(f"⚠️ {date} 环境温度不一致，使用首次输入的 {env_by_date[date]}")
            idx = date_idx[date]
            blocks_by_cat_bucket[cat][idx].extend(expand_blocks(slice_items, PER_LINE_PER_BLOCK))

    # —— 兜底 ——
    left_total = sum(len(v) for v in leftover_by_cat.values())
    if left_total:
        print(f"⚠️ 还有 {left_total} 组未分配。")
        ans = ask("是否给未指定楼层套用【默认日期/数量/温度】？(y=是 / 回车=否→回落到日期分桶)", lower=True)
        if ans == "y":
            default_entries = _prompt_dates_and_limits()
            for cat in CATEGORY_ORDER:
                if not leftover_by_cat.get(cat):
                    continue
                for date, env, slice_items in _distribute_by_dates(leftover_by_cat[cat], default_entries):
                    if not slice_items:
                        continue
                    if date not in date_idx:
                        date_idx[date] = len(buckets)
                        buckets.append({"date": date, "env": env})
                        env_by_date[date] = env
                    elif env_by_date[date] != env:
                        print(f"⚠️ {date} 环境温度不一致，使用首次输入的 {env_by_date[date]}")
                    idx = date_idx[date]
                    blocks_by_cat_bucket[cat][idx].extend(expand_blocks(slice_items, PER_LINE_PER_BLOCK))
                leftover_by_cat[cat] = []
        else:
            grouped_left = {c: leftover_by_cat[c] for c in CATEGORY_ORDER if leftover_by_cat.get(c)}
            if grouped_left:
                buckets2 = prompt_date_buckets(list(grouped_left.keys()))
                later_first = prompt_bucket_priority()
                cat_byb, remain_by_cat = assign_by_buckets(grouped_left, buckets2, later_first)
                ok, auto_last = preview_buckets_generic(cat_byb, remain_by_cat, buckets2, list(grouped_left.keys()))
                if ok:
                    if auto_last:
                        last = len(buckets2) - 1
                        for c in grouped_left.keys():
                            cat_byb[c][last].extend(remain_by_cat[c])
                            remain_by_cat[c] = []
                    blocks_by_cat_bucket2 = expand_blocks_by_bucket(cat_byb)
                    for i, bk in enumerate(buckets2):
                        date, env = bk["date"], bk["env"]
                        if date not in date_idx:
                            date_idx[date] = len(buckets)
                            buckets.append({"date": date, "env": env})
                            env_by_date[date] = env
                        elif env_by_date[date] != env:
                            print(f"⚠️ {date} 环境温度不一致，使用首次输入的 {env_by_date[date]}")
                        idx = date_idx[date]
                        for c in grouped_left.keys():
                            blocks_by_cat_bucket[c][idx].extend(blocks_by_cat_bucket2[c].get(i, []))
                    leftover_by_cat = remain_by_cat
                else:
                    print("❌ 已取消兜底分配。")

    unassigned = sum(len(v) for v in leftover_by_cat.values())

    # —— 日期按升序排序 ——
    order = sorted(range(len(buckets)), key=lambda i: buckets[i]["date"])
    buckets = [buckets[i] for i in order]
    for cat in CATEGORY_ORDER:
        blocks_by_cat_bucket[cat] = {new_i: blocks_by_cat_bucket[cat].get(old_i, []) for new_i, old_i in
                                     enumerate(order)}

    # —— 统一写页 ——
    cats_in_use = [c for c in CATEGORY_ORDER if blocks_by_cat_bucket[c]]
    pages_slices_by_cat = {}
    for cat in cats_in_use:
        blocks_dict = {i: blocks_by_cat_bucket[cat].get(i, []) for i in range(len(buckets))}
        pages_slices_by_cat[cat] = ensure_pages_slices_for_cat(wb, cat, blocks_dict)

    target = make_target_order_generic(pages_slices_by_cat, cats_in_use)
    for idx, name in enumerate(target):
        cur = wb.sheetnames.index(name)
        if cur != idx:
            wb.move_sheet(wb[name], idx - cur)

    total_blocks = 0
    for cat in cats_in_use:
        for i in range(len(buckets)):
            total_blocks += len(blocks_by_cat_bucket[cat].get(i, []))
    prog = Prog(total_blocks, "写入 Excel")
    for i in range(len(buckets)):
        day_pages = []
        for cat in CATEGORY_ORDER:
            if cat not in cats_in_use:
                continue
            pages = pages_slices_by_cat[cat][i]
            blocks = blocks_by_cat_bucket[cat].get(i, [])
            fill_blocks_to_pages(wb, pages, blocks, prog)
            day_pages += pages
        apply_meta_on_pages(wb, day_pages, normalize_date(buckets[i]["date"]), normalize_env(buckets[i]["env"]),
                            auto_instrument=True)
    prog.finish()

    used_names_total = target
    if unassigned:
        print(f"⚠️ 未指派：{unassigned} 组")
    return used_names_total


def try_handle_mode4(mode, wb, grouped, categories_present) -> list | None:
    """模式4兼容钩子。"""
    if mode != "4":
        return None
    return mode4_run(wb, grouped, categories_present)


# ===== 旧法子模式 =====
def prompt_break_submode(has_gz, has_gl):
    """
    交互式选择楼层断点子模式，适配不同数据场景。

    根据是否同时存在钢柱和钢梁提供选项：
    - 同时存在：支持共用断点、分别断点或无断点
    - 仅单类：支持无断点或分别断点
    确保子模式适配实际数据类型。

    Args:
        has_gz: 是否存在钢柱数据（bool）
        has_gl: 是否存在钢梁数据（bool）
    Returns:
        str: 子模式编号（"1"|"2"|"3"）
    """
    if has_gz and has_gl:
        t = ask("断点子模式：1) 柱梁共用断点（简便）  2) 柱梁分别断点  3) 无断点（整单同一天）")
        return t if t in ("1", "2", "3") else "1"
    else:
        t = ask("断点子模式：仅存在单类（或加“其他”） → 3) 无断点  或  2) 分别断点（按各自断点）")
        return t if t in ("2", "3") else "3"


# ===== 主流程 =====
def run_mode(mode: str, wb, grouped, categories_present):
    """按指定模式执行一次导出。"""
    global support_bucket_strategy
    res = try_handle_mode4(mode, wb, grouped, categories_present)
    if res is not None:
        return res

    if mode == "2":
        # —— 旧法：断点 ——
        has_gz = "钢柱" in categories_present
        has_gl = "钢梁" in categories_present
        sub = prompt_break_submode(has_gz, has_gl)

        # 准备 blocks
        blocks_by_cat = {cat: expand_blocks(grouped[cat], PER_LINE_PER_BLOCK) for cat in categories_present}

        if sub == "3":
            # 无断点：按顺序依次排
            pages_by_cat = {}

            def need_pages(lst):
                return math.ceil(len(lst) / BLOCKS_PER_SHEET) if lst else 0

            for cat in categories_present:
                total = need_pages(blocks_by_cat[cat])
                if total == 0:
                    pages_by_cat[cat] = []
                else:
                    if cat == "其他":
                        pages_by_cat[cat] = ensure_total_pages_from(wb, "钢柱", "其他", total)
                    else:
                        pages_by_cat[cat] = ensure_total_pages(wb, cat, total)
            target = []
            for cat in CATEGORY_ORDER:
                if cat in categories_present:
                    target += pages_by_cat[cat]
            for idx, name in enumerate(target):
                cur = wb.sheetnames.index(name)
                if cur != idx: wb.move_sheet(wb[name], idx - cur)

            total_blocks = sum(len(blocks_by_cat[cat]) for cat in categories_present)
            prog = Prog(total_blocks, "写入 Excel")
            for cat in CATEGORY_ORDER:
                if cat in categories_present:
                    fill_blocks_to_pages(wb, pages_by_cat[cat], blocks_by_cat[cat], prog)
            prog.finish()

            d = normalize_date(
                ask("📅 整单日期（20250101 / 2025年1月1日 / 2025 1 1 / 2025.1.1 / 2025-1-1 / 1-1 / 01-01；回车=不写）：") or "")
            e = normalize_env(ask("🌡 整单环境（回车=不写）：") or "")
            apply_meta_on_pages(wb, target, d, e, auto_instrument=True)
            used_names_total = target

        else:
            # 分别断点（若同时有柱&梁，可选择共用；“其他”总是用自己的断点）
            same_breaks = None
            if has_gz and has_gl and sub == "1":
                same_breaks = prompt_floor_breaks("钢柱/钢梁（共用）")
            breaks_by_cat = {}
            for cat in categories_present:
                if cat == "支撑":
                    prompt_support_strategy_for_bucket()
                    if support_bucket_strategy == "floor":
                        breaks_by_cat[cat] = prompt_floor_breaks(cat)
                    else:
                        breaks_by_cat[cat] = []  # 支撑不做断点分段
                elif cat in ("钢柱", "钢梁") and same_breaks is not None:
                    breaks_by_cat[cat] = same_breaks
                else:
                    breaks_by_cat[cat] = prompt_floor_breaks(cat)

            # 分段
            byseg = {cat: defaultdict(list) for cat in categories_present}
            for cat in categories_present:
                if cat == "支撑" and support_bucket_strategy != "floor":
                    byseg[cat][0] = blocks_by_cat[cat]
                else:
                    for b in blocks_by_cat[cat]:
                        seg = segment_index(floor_of(b["name"]), breaks_by_cat[cat])
                        byseg[cat][seg].append(b)
            rounds = max((max(byseg[cat].keys()) if byseg[cat] else 0) for cat in categories_present) + 1

            # 预分配页
            def pages_needed(lst):
                return math.ceil(len(lst) / BLOCKS_PER_SHEET) if lst else 0

            pages_pool_by_cat = {}
            for cat in categories_present:
                total_pages = sum(pages_needed(byseg[cat].get(i, [])) for i in range(rounds))
                if total_pages == 0:
                    pages_pool_by_cat[cat] = []
                else:
                    if cat == "其他":
                        pages_pool_by_cat[cat] = ensure_total_pages_from(wb, "钢柱", "其他", total_pages)
                    else:
                        pages_pool_by_cat[cat] = ensure_total_pages(wb, cat, total_pages)

            # 计算最终顺序：按轮次交错（柱→梁→支撑→其他）
            target = []
            cursor = {cat: 0 for cat in categories_present}
            for i in range(rounds):
                for cat in CATEGORY_ORDER:
                    if cat not in categories_present: continue
                    need = pages_needed(byseg[cat].get(i, []))
                    pool = pages_pool_by_cat[cat]
                    target += pool[cursor[cat]:cursor[cat] + need]
                    cursor[cat] += need

            # 排序成最终顺序
            for idx, name in enumerate(target):
                cur = wb.sheetnames.index(name)
                if cur != idx: wb.move_sheet(wb[name], idx - cur)

            # 写入（带进度）
            total_blocks = sum(len(byseg[cat].get(i, [])) for cat in categories_present for i in range(rounds))
            prog = Prog(total_blocks, "写入 Excel")
            cursor = {cat: 0 for cat in categories_present}
            for i in range(rounds):
                for cat in CATEGORY_ORDER:
                    if cat not in categories_present: continue
                    seg_blocks = byseg[cat].get(i, [])
                    need = pages_needed(seg_blocks)
                    pool = pages_pool_by_cat[cat]
                    fill_blocks_to_pages(wb, pool[cursor[cat]:cursor[cat] + need], seg_blocks, prog)
                    cursor[cat] += need
            prog.finish()

            # 断点法：整单不分日期，仪器按页自动识别
            apply_meta_on_pages(wb, target, "", "", auto_instrument=True)
            used_names_total = target

    elif mode == "3":
        # —— 简单模式：一次日期/温度；不分段；按 CATEGORY_ORDER 排 ——
        blocks_by_cat = {cat: expand_blocks(grouped[cat], PER_LINE_PER_BLOCK) for cat in categories_present}
        pages_by_cat = {}

        def need_pages(lst):
            return math.ceil(len(lst) / BLOCKS_PER_SHEET) if lst else 0

        for cat in categories_present:
            total = need_pages(blocks_by_cat[cat])
            if total == 0:
                pages_by_cat[cat] = []
            else:
                if cat == "其他":
                    pages_by_cat[cat] = ensure_total_pages_from(wb, "钢柱", "其他", total)
                else:
                    pages_by_cat[cat] = ensure_total_pages(wb, cat, total)
        target = []
        for cat in CATEGORY_ORDER:
            if cat in categories_present:
                target += pages_by_cat[cat]
        for idx, name in enumerate(target):
            cur = wb.sheetnames.index(name)
            if cur != idx: wb.move_sheet(wb[name], idx - cur)

        total_blocks = sum(len(blocks_by_cat[cat]) for cat in categories_present)
        prog = Prog(total_blocks, "写入 Excel")
        for cat in CATEGORY_ORDER:
            if cat in categories_present:
                fill_blocks_to_pages(wb, pages_by_cat[cat], blocks_by_cat[cat], prog)
        prog.finish()

        d = normalize_date(
            ask("📅 日期：20250101 / 2025年1月1日 / 2025 1 1 / 2025.1.1 / 2025-1-1 / 1-1 / 01-01；（回车=不写）：") or "")
        e = normalize_env(ask("🌡 环境温度（回车=不写）：") or "")
        apply_meta_on_pages(wb, target, d, e, auto_instrument=True)
        used_names_total = target

    else:
        # —— 新法：日期分桶（泛化） ——
        buckets = prompt_date_buckets(categories_present)
        later_first = prompt_bucket_priority()  # 回车=是
        cat_byb, remain_by_cat = assign_by_buckets(grouped, buckets, later_first)
        ok, auto_last = preview_buckets_generic(cat_byb, remain_by_cat, buckets, categories_present)
        if not ok:
            print("已取消。");
            return
        if auto_last:
            last = len(buckets) - 1
            for cat in categories_present:
                cat_byb[cat][last].extend(remain_by_cat[cat])
                remain_by_cat[cat] = []

        blocks_by_cat_bucket = expand_blocks_by_bucket(cat_byb)

        # 为每个类别生成每天的页切片
        pages_slices_by_cat = {}
        for cat in categories_present:
            pages_slices_by_cat[cat] = ensure_pages_slices_for_cat(wb, cat, blocks_by_cat_bucket[cat])

        # 最终顺序：按轮次交错（柱→梁→支撑→其他）
        target = make_target_order_generic(pages_slices_by_cat, categories_present)
        for idx, name in enumerate(target):
            cur = wb.sheetnames.index(name)
            if cur != idx: wb.move_sheet(wb[name], idx - cur)

        # 写入（带进度）
        total_blocks = 0
        for cat in categories_present:
            total_blocks += sum(len(v) for v in blocks_by_cat_bucket[cat].values())
        prog = Prog(total_blocks, "写入 Excel")

        for i in range(len(buckets)):
            # 逐天写
            day_pages = []
            for cat in CATEGORY_ORDER:
                if cat in categories_present:
                    fill_blocks_to_pages(wb, pages_slices_by_cat[cat][i], blocks_by_cat_bucket[cat].get(i, []), prog)
                    day_pages += pages_slices_by_cat[cat][i]
            # 日期/环境/仪器
            apply_meta_on_pages(wb, day_pages, buckets[i]["date"], buckets[i]["env"], auto_instrument=True)
        prog.finish()

        used_names_total = target

    return used_names_total

    # ===== 预处理与模式运行封装 =====


def prepare_from_word(src: Path):
    groups_all_tables, all_rows = read_groups_from_doc(src)
    grouped = defaultdict(list)
    for g in groups_all_tables:
        grouped[kind_of(g["name"])].append(g)
    categories_present = [cat for cat in CATEGORY_ORDER if grouped.get(cat)]
    print("📊 识别： " + "、".join(f"{cat} {len(grouped.get(cat, []))}" for cat in categories_present))

    doc_out = build_summary_doc_with_progress(all_rows)
    set_doc_font_progress(doc_out, DEFAULT_FONT_PT)
    out_docx = src.with_name("汇总原始记录.docx")
    print("💾 正在保存汇总 Word …")

    save_docx_safe(doc_out, out_docx)
    print(f"✅ 汇总 Word 已保存：{out_docx}")
    return grouped, categories_present


def run_with_mode(src: Path, grouped, categories_present, meta):
    tpl_path = XLSX_WITH_SUPPORT_DEFAULT  # 始终使用有支撑模板
    if not tpl_path.exists():
        raise FileNotFoundError(f"Excel 模板不存在：{tpl_path}")

    wb = load_workbook_safe(tpl_path)

    try:
        mode = prompt_mode()
        used_names_total = run_mode(mode, wb, grouped, categories_present)
    except BackStep:
        return

    apply_meta_fixed(wb, categories_present, meta)
    enforce_mu_font(wb)
    cleanup_unused_sheets(wb, used_names_total, bases=tuple(CATEGORY_ORDER))

    def unique_out_path(dest_dir: Path, stem: str) -> Path:
        cand = dest_dir / f"{stem}.xlsx"
        if not cand.exists():
            return cand
        i = 1
        while True:
            cand = dest_dir / f"{stem}({i}).xlsx"
            if not cand.exists():
                return cand
            i += 1

    final_path = unique_out_path(src.parent, f"{TITLE}_报告版")
    save_workbook_safe(wb, final_path)
    print(f"✅ Excel 已保存：{final_path}")
    print("✔ 完成。本次导出结束。")
    # 只在本进程第一次成功导出后，给个低调彩蛋提示
    global _hint_shown
    if not _hint_shown:
        print(dark_hint("Maybe you can try entering 'k' the next time you input the file path."))
        _hint_shown = True


# ===== 顶层交互循环 =====
def main():
    print(f" {TITLE} — {VERSION}")
    while True:
        path = ask_path()
        if path is None:
            continue
        if path == "__QUIT__":
            print("Bye")
            break
        if not is_valid_path(path):
            print("× 路径无效。")
            continue
        try:
            src = Path(path)
            print(f"✅ 使用 Word：{src}")
            global support_bucket_strategy
            support_bucket_strategy = None

            grouped, categories_present = prepare_from_word(src)

            proj = ask("工程名称（回车可空）：")
            order = ask("委托编号（回车可空）：")
            meta = {"proj": proj or "", "order": order or ""}

            run_with_mode(src, grouped, categories_present, meta)

        except FileInUse as e:
            # ↓↓↓ 友好提示，不打印堆栈，不吓用户
            print("\n⚠️  文件被占用，无法读写：")
            print(f"   - {e}")
            print("✅  请关闭相关的 Excel / Word / 预览窗口（含资源管理器预览窗格），然后重新运行本程序。\n")
            # 直接回到主循环
            continue

        except Exception as e:
            # 其他异常仍提示，但不长篇堆栈
            print(f"× 出错：{e}")
            continue


# ===== 读取 Word 分组 =====
def read_groups_from_doc(path: Path):
    """
    从Word文档中读取并解析构件数据组，返回结构化分组数据和原始行数据。

    流程：
    1. 打开Word文档并遍历所有表格，筛选含“测点1”和“平均值”的有效数据表格
    2. 对每个有效表格提取数据行（带进度提示）
    3. 将提取的原始行数据转换为按构件名称分组的结构化数据

    结构化数据组包含构件名称和对应的测点数据（8个读数+1个平均值），适配后续Excel填充需求。

    Args:
        path: Word文档路径（Path对象）
    Returns:
        tuple: 包含两个元素的元组：
            - 构件数据组列表（list[dict]），每个元素含'name'（构件名）和'data'（数据行列表）
            - 所有原始数据行列表（list[dict]），含提取的测点值、平均值等原始信息
    """
    doc = Document(str(path))
    all_rows = []
    tables = doc.tables
    T = sum(1 for t in tables if is_data_table(t))  # noqa
    used = 0
    for tbl in tables:
        if not is_data_table(tbl):
            continue
        used += 1
        part = extract_rows_with_progress(tbl, used, T)
        if part: all_rows.extend(part)
    return groups_from_your_rows(all_rows), all_rows


if __name__ == "__main__":
    main()

    # v4.2.3
